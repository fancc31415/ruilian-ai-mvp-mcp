"""
匹配报告生成模块。
- generate_match_reasons: 给Top N机构逐条生成『为什么匹配 + 怎么谈』（LLM优先，模板降级，带缓存）
- export_markdown / export_docx / export_pdf: 三种格式导出
"""
import hashlib
import json
from datetime import datetime

REASON_SYSTEM_PROMPT = """你是一名资深FA顾问，正在为一份BP撰写"机构匹配理由与沟通建议"。
输入是企业信息和一家候选投资机构的信息，请输出JSON：
{
  "match_reason": "为什么这家机构值得联系，2-3句话，结合赛道/阶段/规模/机构风格具体展开，不要套话",
  "talking_points": ["与该机构沟通时建议突出的1-2个点，要具体，结合企业亮点和机构偏好"],
  "caution": "如果有需要注意的匹配风险或需要提前准备的材料，用一句话说明；没有则填'无明显风险点'"
}
只返回JSON，不要多余文字。
"""


def _cache_key(bp: dict, inst: dict) -> str:
    """用BP内容+机构名算一个稳定的哈希key。BP哪怕字段顺序不同、只要内容一样，key也一样，
    这样"改了个不影响匹配理由的字段（比如公司名打错重传）"不会误命中旧缓存里过时的内容——
    只取真正影响匹配理由生成的字段参与哈希，其余变动允许命中缓存。"""
    relevant_bp = {
        "sectors": bp.get("sectors"),
        "stage": bp.get("stage"),
        "funding_ask_wan": bp.get("funding_ask_wan"),
        "business_summary": bp.get("business_summary"),
        "highlights": bp.get("highlights"),
    }
    raw = json.dumps(relevant_bp, sort_keys=True, ensure_ascii=False) + "||" + inst.get("name", "")
    return hashlib.md5(raw.encode("utf-8")).hexdigest()


def generate_match_reasons(bp: dict, matched_institutions: list, llm_client, cache: dict = None) -> list:
    """cache: 外部传入的dict（比如st.session_state里的一个字典），跨多次点击复用。
    不传的话退化成一次性生成，不缓存（比如给MCP服务器这种无状态调用场景用）。"""
    if cache is None:
        cache = {}

    enriched = []
    for inst in matched_institutions:
        key = _cache_key(bp, inst)
        if key in cache:
            reason_data = cache[key]
        else:
            if llm_client.available:
                try:
                    user_prompt = (
                        f"企业信息：{bp}\n\n候选机构信息：{ {k: v for k, v in inst.items() if k not in ['score_breakdown']} }"
                    )
                    reason_data = llm_client.chat_json(REASON_SYSTEM_PROMPT, user_prompt)
                except Exception:
                    reason_data = _template_reason(bp, inst)
            else:
                reason_data = _template_reason(bp, inst)
            cache[key] = reason_data

        enriched.append({**inst, **reason_data})
    return enriched


def _template_reason(bp: dict, inst: dict) -> dict:
    """没有LLM时的模板化降级方案，保证报告字段完整可用。"""
    matched = "、".join(inst.get("matched_sectors", [])) or "赛道信息待人工核对"
    return {
        "match_reason": (
            f"{inst['name']} 在{matched}赛道有布局，投资阶段覆盖{'/'.join(inst.get('stages', []))}，"
            f"单笔规模区间{inst.get('check_size_min_wan')}万-{inst.get('check_size_max_wan')}万元，"
            f"与本项目融资需求综合匹配度较高。"
        ),
        "talking_points": [f"强调项目在{matched}赛道的差异化优势，对齐机构历史投资偏好"],
        "caution": "此为模板化生成（未配置LLM API Key），建议人工顾问复核后再对外沟通。",
    }


def export_markdown(bp: dict, enriched_matches: list) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [
        f"# 睿链AI · 融资匹配报告",
        f"生成时间：{now}",
        "",
        "## 一、项目概览",
        f"- **公司名称**：{bp.get('company_name', '未提供')}",
        f"- **所属赛道**：{'、'.join(bp.get('sectors', []))}",
        f"- **融资阶段**：{bp.get('stage', '未提供')}",
        f"- **本轮融资金额**：{bp.get('funding_ask_wan', 0)} 万元",
        f"- **估值**：{bp.get('valuation_wan', 0)} 万元" if bp.get("valuation_wan") else "- **估值**：未提供",
        f"- **商业模式概括**：{bp.get('business_summary', '未提供')}",
        f"- **核心亮点**：{'；'.join(bp.get('highlights', [])) or '未提供'}",
        f"- **需人工核对项**：{'；'.join(bp.get('risks_or_gaps', [])) or '无'}",
        "",
        f"## 二、匹配机构清单（Top {len(enriched_matches)}）",
        "",
    ]

    for i, m in enumerate(enriched_matches, 1):
        lines += [
            f"### {i}. {m['name']}  ·  综合匹配度 {m['match_score']}分",
            f"- **类型**：{m.get('type', '')}　**代表案例**：{m.get('notable', '')}",
            f"- **赛道覆盖**：{'、'.join(m.get('sectors', []))}",
            f"- **投资阶段**：{'、'.join(m.get('stages', []))}",
            f"- **单笔规模**：{m.get('check_size_min_wan')}万 - {m.get('check_size_max_wan')}万元",
            f"- **打分明细**：赛道{m['score_breakdown']['赛道匹配']} / 阶段{m['score_breakdown']['阶段匹配']} "
            f"/ 规模{m['score_breakdown']['规模匹配']} / 其他{m['score_breakdown']['其他信号']}",
            f"- **匹配理由**：{m.get('match_reason', '')}",
            f"- **沟通建议**：{'；'.join(m.get('talking_points', []))}",
            f"- **注意事项**：{m.get('caution', '')}",
            "",
        ]

    lines.append("---")
    lines.append("*本报告由睿链AI自动生成，匹配结果仅供参考，最终决策建议结合人工顾问复核意见。*")
    return "\n".join(lines)


def export_docx(bp: dict, enriched_matches: list) -> bytes:
    """导出Word文档，返回文件的二进制内容（配合st.download_button直接用）。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    import io

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc = Document()

    title = doc.add_heading("睿链AI · 融资匹配报告", level=0)
    doc.add_paragraph(f"生成时间：{now}").italic = True

    doc.add_heading("一、项目概览", level=1)
    overview_fields = [
        ("公司名称", bp.get("company_name", "未提供")),
        ("所属赛道", "、".join(bp.get("sectors", []))),
        ("融资阶段", bp.get("stage", "未提供")),
        ("本轮融资金额", f"{bp.get('funding_ask_wan', 0)} 万元"),
        ("估值", f"{bp.get('valuation_wan', 0)} 万元" if bp.get("valuation_wan") else "未提供"),
        ("商业模式概括", bp.get("business_summary", "未提供")),
        ("核心亮点", "；".join(bp.get("highlights", [])) or "未提供"),
        ("需人工核对项", "；".join(bp.get("risks_or_gaps", [])) or "无"),
    ]
    for label, value in overview_fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(str(value))

    doc.add_heading(f"二、匹配机构清单（Top {len(enriched_matches)}）", level=1)
    for i, m in enumerate(enriched_matches, 1):
        doc.add_heading(f"{i}. {m['name']}　·　综合匹配度 {m['match_score']}分", level=2)
        bd = m["score_breakdown"]
        detail_fields = [
            ("类型", m.get("type", "")),
            ("代表案例", m.get("notable", "")),
            ("赛道覆盖", "、".join(m.get("sectors", []))),
            ("投资阶段", "、".join(m.get("stages", []))),
            ("单笔规模", f"{m.get('check_size_min_wan')}万 - {m.get('check_size_max_wan')}万元"),
            ("打分明细", f"赛道{bd['赛道匹配']} / 阶段{bd['阶段匹配']} / 规模{bd['规模匹配']} / 其他{bd['其他信号']}"),
            ("匹配理由", m.get("match_reason", "")),
            ("沟通建议", "；".join(m.get("talking_points", []))),
            ("注意事项", m.get("caution", "")),
        ]
        for label, value in detail_fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}：").bold = True
            p.add_run(str(value))

    footer = doc.add_paragraph()
    footer_run = footer.add_run("本报告由睿链AI自动生成，匹配结果仅供参考，最终决策建议结合人工顾问复核意见。")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(bp: dict, enriched_matches: list) -> bytes:
    """导出PDF，返回文件的二进制内容。用纯Python的fpdf2 + 自带的中文字体，
    不依赖系统层面的PDF渲染引擎，避免在Streamlit Cloud这类托管环境上因为缺系统依赖而部署失败。"""
    from fpdf import FPDF
    import os

    font_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts", "wqy-microhei.ttf")

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("wqy", "", font_path)
    pdf.set_font("wqy", size=16)
    pdf.cell(0, 12, "睿链AI · 融资匹配报告", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("wqy", size=9)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    pdf.cell(0, 8, f"生成时间：{now}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    def add_heading(text, size=13):
        pdf.set_font("wqy", size=size)
        pdf.multi_cell(0, 9, text)
        pdf.ln(1)

    def add_field(label, value, indent=4):
        pdf.set_font("wqy", size=10)
        pdf.set_x(pdf.l_margin + indent)
        pdf.multi_cell(0, 6.5, f"{label}：{value}")

    add_heading("一、项目概览")
    overview_fields = [
        ("公司名称", bp.get("company_name", "未提供")),
        ("所属赛道", "、".join(bp.get("sectors", []))),
        ("融资阶段", bp.get("stage", "未提供")),
        ("本轮融资金额", f"{bp.get('funding_ask_wan', 0)} 万元"),
        ("估值", f"{bp.get('valuation_wan', 0)} 万元" if bp.get("valuation_wan") else "未提供"),
        ("商业模式概括", bp.get("business_summary", "未提供")),
        ("核心亮点", "；".join(bp.get("highlights", [])) or "未提供"),
    ]
    for label, value in overview_fields:
        add_field(label, value)
    pdf.ln(3)

    add_heading(f"二、匹配机构清单（Top {len(enriched_matches)}）")
    for i, m in enumerate(enriched_matches, 1):
        add_heading(f"{i}. {m['name']}  ·  匹配度 {m['match_score']}分", size=11)
        bd = m["score_breakdown"]
        detail_fields = [
            ("类型", m.get("type", "")),
            ("赛道覆盖", "、".join(m.get("sectors", []))),
            ("投资阶段", "、".join(m.get("stages", []))),
            ("单笔规模", f"{m.get('check_size_min_wan')}万 - {m.get('check_size_max_wan')}万元"),
            ("匹配理由", m.get("match_reason", "")),
            ("沟通建议", "；".join(m.get("talking_points", []))),
        ]
        for label, value in detail_fields:
            add_field(label, value)
        pdf.ln(2)

    pdf.set_font("wqy", size=8)
    pdf.multi_cell(0, 5, "本报告由睿链AI自动生成，匹配结果仅供参考，最终决策建议结合人工顾问复核意见。")

    return bytes(pdf.output())
